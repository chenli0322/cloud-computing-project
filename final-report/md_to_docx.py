"""
Converts the merged Part-1 / Part-2 weekly progress reports (Markdown) into
DOCX files that can be uploaded directly to Brightspace.

Run from final-report/:
    python md_to_docx.py

Outputs:
    Part1_Weekly_Progress_Reports.docx
    Part2_Weekly_Progress_Reports.docx
"""
from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent

INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_inline(p, text: str, *, base_size: int = 11):
    """Add a paragraph run, handling inline code/bold/links sequentially."""
    # Find all spans (link / bold / code) and split text accordingly.
    pieces = []
    i = 0
    while i < len(text):
        # Find next markdown span
        m_link = LINK_RE.search(text, i)
        m_bold = BOLD_RE.search(text, i)
        m_code = INLINE_CODE_RE.search(text, i)
        candidates = [m for m in (m_link, m_bold, m_code) if m is not None]
        if not candidates:
            pieces.append(("plain", text[i:]))
            break
        m = min(candidates, key=lambda x: x.start())
        if m.start() > i:
            pieces.append(("plain", text[i:m.start()]))
        if m is m_link:
            pieces.append(("link", m.group(1), m.group(2)))
        elif m is m_bold:
            pieces.append(("bold", m.group(1)))
        else:
            pieces.append(("code", m.group(1)))
        i = m.end()

    for piece in pieces:
        kind = piece[0]
        if kind == "plain":
            r = p.add_run(piece[1])
            r.font.size = Pt(base_size)
        elif kind == "bold":
            r = p.add_run(piece[1])
            r.bold = True
            r.font.size = Pt(base_size)
        elif kind == "code":
            r = p.add_run(piece[1])
            r.font.name = "Consolas"
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif kind == "link":
            r = p.add_run(piece[1])
            r.font.size = Pt(base_size)
            r.font.color.rgb = RGBColor(0x06, 0x5F, 0xD8)
            r.font.underline = True


def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code_block = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        for line in code_buf:
            r = p.add_run(line + "\n")
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        code_buf = []

    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # Filter out separator rows (e.g. |---|---|)
        rows = [r for r in table_rows if not all(cell.strip().replace("-", "").replace(":", "") == "" for cell in r)]
        if not rows:
            in_table = False
            table_rows = []
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Light Grid Accent 1"
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                cell_obj = table.rows[ri].cells[ci]
                cell_obj.text = ""
                p = cell_obj.paragraphs[0]
                add_inline(p, cell.strip(), base_size=10)
                if ri == 0:
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()  # spacer
        in_table = False
        table_rows = []

    for raw in lines:
        line = raw.rstrip()

        # Code block toggle
        if line.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            continue
        if in_code_block:
            code_buf.append(line)
            continue

        # Tables
        if line.startswith("|") and line.endswith("|"):
            flush_code()
            cells = [c for c in line.strip("|").split("|")]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        else:
            if in_table:
                flush_table()

        if not line.strip():
            doc.add_paragraph()
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            continue

        # Horizontal rule
        if set(line.strip()) <= {"-"} and len(line.strip()) >= 3:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run("─" * 80)
            r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            continue

        # Blockquotes
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, line[2:].strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            continue

        # Bulleted list
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            continue
        # Indented bullet
        if line.startswith("  - ") or line.startswith("    - "):
            indent = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.5 + 0.25 * (indent - 1))
            add_inline(p, line.lstrip()[2:])
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, line)

    flush_code()
    flush_table()

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


def main():
    for part in (1, 2):
        src = HERE / f"Part{part}_Weekly_Progress_Reports.md"
        dst = HERE / f"Part{part}_Weekly_Progress_Reports.docx"
        md_to_docx(src, dst)


if __name__ == "__main__":
    main()
